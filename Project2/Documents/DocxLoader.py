from dataclasses import dataclass, field
from typing import Dict, Any, Optional, List

@dataclass
class Document:
    """
    A universal data structure for all document types.
    Works with loaders, chunkers, embeddings, and LLMs.
    """
    page_content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    blob: Optional[bytes] = None       # raw binary (for images)
    resources: Optional[List[str]] = field(default_factory=list)  # extracted file paths, images, etc.
    id: Optional[str] = None
    
    def __repr__(self):
        preview = self.page_content[:80].replace("\n", " ") + ("..." if len(self.page_content) > 80 else "")
        return f"<Document id={self.id} source={self.metadata.get('source','')} content='{preview}'>"
from docx import Document as DocxDocument


class LoadDOCX:
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> list:
        docx = DocxDocument(self.file_path)
        documents = []

        for i, para in enumerate(docx.paragraphs):
            text = para.text.strip()
            if text:
                documents.append(Document(
                    page_content=text,
                    metadata={"source": self.file_path, "paragraph": i + 1}
                ))
        self._documents = documents
        return documents
    def __iter__(self):
        if self._documents is None:
            self.load()
        return iter(self._documents)
